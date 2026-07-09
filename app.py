import os
import io
import base64
import tempfile
import subprocess
import unicodedata
import json
import urllib.request
import urllib.parse
import urllib.error
from flask import Flask, request, render_template, jsonify
from werkzeug.utils import secure_filename
import pytesseract
from PIL import Image, ImageOps
from pdf2image import convert_from_bytes
from docx import Document
from docx.oxml import OxmlElement
import copy
import pypdf
from deep_translator import GoogleTranslator

# Import your converter and utils
from converter import PunjabiFontConverter
from utils import (
    sanitize_unicode, is_gurmukhi_unicode, detect_legacy_font,
    FONT_FAMILY_MAP, DROPDOWN_TO_CONVERTER, CONV_SOURCE_UNICODE,
    split_gurmukhi_segments, WORD_FONT_MAP, LEGACY_FONTS
)

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 32 * 1024 * 1024  # 32 MB

conv = PunjabiFontConverter()

# ---------- Poppler path detection (same as desktop) ----------
def find_poppler():
    paths = [
        os.environ.get('POPPLER_PATH', ''),
        r'C:\poppler\Library\bin',
        r'C:\poppler\bin',
        r'C:\Program Files\poppler\bin',
        r'C:\Program Files\poppler\Library\bin',
    ]
    for p in paths:
        if p and os.path.exists(p):
            return p
    try:
        result = subprocess.run(['where', 'pdftoppm'], capture_output=True, text=True, shell=True)
        if result.returncode == 0:
            return os.path.dirname(result.stdout.splitlines()[0].strip())
    except:
        pass
    return None

POPPLER_PATH = find_poppler()
if POPPLER_PATH:
    print(f"✅ Poppler: {POPPLER_PATH}")
else:
    print("ℹ️ Local Poppler not found (Gemini Cloud OCR will be used for PDF files).")

# ---------- Tesseract path detection ----------
def find_tesseract():
    try:
        subprocess.run(['tesseract', '--version'], capture_output=True, check=True)
        return 'tesseract'
    except:
        pass
    candidates = [
        r'C:\Program Files\Tesseract-OCR\tesseract.exe',
        r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
    ]
    for c in candidates:
        if os.path.exists(c):
            return c
    return None

TESSERACT_PATH = find_tesseract()
if TESSERACT_PATH:
    print(f"✅ Tesseract: {TESSERACT_PATH}")
    if TESSERACT_PATH != 'tesseract':
        pytesseract.pytesseract.tesseract_cmd = TESSERACT_PATH
else:
    print("ℹ️ Local Tesseract not found (Gemini Cloud OCR will be used for Image files).")

# ---------- Gemini API OCR Function ----------
def ocr_via_gemini(file_bytes: bytes, mime_type: str, api_key: str) -> str:
    """Call Gemini 2.5 Flash API to extract text from image or PDF, preserving format & layout."""
    url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash:generateContent?key={api_key}"
    b64_data = base64.b64encode(file_bytes).decode('utf-8')
    
    prompt = (
        "You are an expert Punjabi OCR engine specializing in legal and police documents. "
        "Extract all text (both typed and handwritten) from the provided document in its original languages (Gurmukhi Punjabi, English, numbers). "
        "You must preserve the original layout, columns, alignment, spacing, and formatting as closely as possible. "
        "Carefully transcribe all handwritten notes, inspector marks, signatures, dates, and names at the bottom exactly where they appear in the original layout. "
        "Do not omit, summarize, or translate anything. Output only the extracted text without any introduction or notes."
    )
    
    payload = {
        "contents": [
            {
                "parts": [
                    {"text": prompt},
                    {
                        "inlineData": {
                            "mimeType": mime_type,
                            "data": b64_data
                        }
                    }
                ]
            }
        ]
    }
    
    req_data = json.dumps(payload).encode('utf-8')
    req = urllib.request.Request(
        url,
        data=req_data,
        headers={'Content-Type': 'application/json'}
    )
    
    try:
        with urllib.request.urlopen(req, timeout=30) as response:
            res_body = response.read().decode('utf-8')
            res_json = json.loads(res_body)
            text = res_json['candidates'][0]['content']['parts'][0]['text']
            return text
    except urllib.error.HTTPError as e:
        error_msg = e.read().decode('utf-8')
        try:
            err_json = json.loads(error_msg)
            message = err_json.get('error', {}).get('message', str(e))
        except:
            message = error_msg or str(e)
        raise Exception(f"Gemini API Error: {message}")
    except Exception as e:
        raise Exception(f"Failed to connect to Gemini: {str(e)}")

# ---------- OCR functions (exactly as in desktop app) ----------
def ocr_image_from_pdf(img: Image.Image) -> str:
    # Use pan+eng+hin to read multilingual text
    config = '--oem 3 --psm 4 -l pan+eng+hin'
    processed = ImageOps.grayscale(img).point(lambda x: 0 if x < 140 else 255, '1')
    return pytesseract.image_to_string(processed, config=config)

def ocr_image(img: Image.Image) -> str:
    # Use pan+eng+hin to read multilingual text
    config = '--oem 3 --psm 4 -l pan+eng+hin'
    processed = ImageOps.grayscale(img).point(lambda x: 0 if x < 140 else 255, '1')
    return pytesseract.image_to_string(processed, config=config)

# ---------- Routes ----------
@app.route('/')
def index():
    return render_template('index.html',
                           font_families=list(FONT_FAMILY_MAP.keys()),
                           DROPDOWN_TO_CONVERTER=DROPDOWN_TO_CONVERTER)

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    if not file.filename.lower().endswith(('.pdf', '.docx', '.png', '.jpg', '.jpeg')):
        return jsonify({'error': 'File type not allowed'}), 400

    ext = file.filename.rsplit('.', 1)[1].lower()
    file_bytes = file.read()
    filename = secure_filename(file.filename)

    # Check for Gemini API key in custom header or request body
    gemini_key = request.headers.get('X-Gemini-API-Key') or request.form.get('api_key')

    try:
        # Advanced OCR with Gemini (if API Key provided and file is PDF/Image)
        if gemini_key and ext in ('pdf', 'png', 'jpg', 'jpeg'):
            mime_type = 'application/pdf' if ext == 'pdf' else f'image/{ext if ext != "jpg" else "jpeg"}'
            try:
                extracted_text = ocr_via_gemini(file_bytes, mime_type, gemini_key)
                return jsonify({
                    'text': sanitize_unicode(extracted_text),
                    'file_type': ext,
                    'ocr_engine': 'gemini'
                })
            except Exception as gemini_err:
                print(f"⚠️ Gemini OCR failed, falling back to local OCR: {gemini_err}")
                # continue to local fallback

        if ext == 'docx':
            doc = Document(io.BytesIO(file_bytes))
            full_text = '\n'.join(p.text for p in doc.paragraphs)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text += '\n' + cell.text
            b64 = base64.b64encode(file_bytes).decode('utf-8')
            return jsonify({
                'text': sanitize_unicode(full_text),
                'file_type': 'docx',
                'file_b64': b64,
                'filename': filename,
                'ocr_engine': 'local'
            })

        elif ext == 'pdf':
            # Try pypdf extraction first
            text = ""
            try:
                reader = pypdf.PdfReader(io.BytesIO(file_bytes))
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text
            except:
                pass

            if text and is_gurmukhi_unicode(text):
                return jsonify({'text': sanitize_unicode(text), 'file_type': 'pdf', 'ocr_engine': 'local'})

            # OCR fallback
            if POPPLER_PATH is None:
                return jsonify({'error': 'Poppler not installed. Local PDF parser cannot run OCR.'}), 500

            pages = convert_from_bytes(file_bytes, dpi=250, poppler_path=POPPLER_PATH)
            ocr_text = "\n".join(ocr_image_from_pdf(p) for p in pages)
            text = ocr_text

            if is_gurmukhi_unicode(text):
                final = text
            else:
                detected_font = detect_legacy_font(text)
                if detected_font and detected_font != "Unicode":
                    final = conv.convert_font(text, detected_font, "Unicode")
                else:
                    final = text
            return jsonify({'text': sanitize_unicode(final), 'file_type': 'pdf', 'ocr_engine': 'local'})

        elif ext in ('png', 'jpg', 'jpeg'):
            img = Image.open(io.BytesIO(file_bytes))
            text = ocr_image(img)
            if is_gurmukhi_unicode(text):
                final = text
            else:
                detected_font = detect_legacy_font(text)
                if detected_font and detected_font != "Unicode":
                    final = conv.convert_font(text, detected_font, "Unicode")
                else:
                    final = text
            return jsonify({'text': sanitize_unicode(final), 'file_type': 'image', 'ocr_engine': 'local'})

    except Exception as e:
        err_msg = str(e)
        if "tesseract" in err_msg.lower() or "poppler" in err_msg.lower() or "no such file" in err_msg.lower():
            err_msg = "Local OCR engine is not available on this cloud server. Please configure your Free Gemini API Key in Settings (top right) to enable high-quality visual OCR."
        return jsonify({'error': err_msg}), 500

def convert_mixed_legacy_to_unicode(text, from_font, conv):
    """
    Convert legacy Gurmukhi text to Unicode while keeping English words unchanged.
    """
    # A set of common English words, including template terms from FIR documents
    english_words = {
        "first", "information", "report", "under", "section", "district", "sadar", "ferozepur", 
        "fir", "date", "time", "hrs", "acts", "narcotic", "drugs", "and", "psychotropic", "substances", 
        "act", "occurrence", "of", "offence", "day", "sunday", "from", "to", "period", "pahar", 
        "received", "general", "diary", "reference", "entry", "type", "place", "direction", 
        "distance", "ps", "south", "km", "beat", "address", "in", "case", "outside", "limit", 
        "police", "station", "then", "name", "state", "complainant", "informant", "asi", "no", 
        "father", "mother", "husband", "year", "birth", "nationality", "india", "uid", "passport", 
        "issue", "id", "details", "ration", "card", "voter", "driving", "license", "pan", "number", 
        "occupation", "present", "unknown", "punjab", "permanent", "phone", "mobile", "known", 
        "suspected", "accused", "full", "particulars", "more", "than", "alias", "relative", 
        "khai", "pheme", "ke", "reasons", "delay", "reporting", "necessary", "properties", "interest", 
        "property", "category", "value", "rs", "stolen", "inquest", "ud", "any", "contents", "phg", 
        "ndps", "action", "taken", "reveals", "commission", "directed", "rank", "assistant", "sub", 
        "investigation", "refused", "due", "transferred", "jurisdiction", "read", "over", "admitted", 
        "correctly", "recorded", "copy", "given", "free", "signature", "thumb", "impression", 
        "officer", "charge", "inspector", "court", "dispatch", "attachment", "item", "physical", 
        "features", "deformities", "sex", "male", "height", "cms", "complexion", "identification", 
        "mark", "peculiarities", "teeth", "hair", "eyes", "habit", "dress", "leucoderm", "scar", 
        "tattoo", "fields", "entered", "gives", "one", "particulars", "suspect", "written", "the", 
        "a", "an", "is", "are", "was", "were", "be", "been", "have", "has", "had", "do", "does", 
        "did", "but", "or", "as", "if", "when", "while", "at", "by", "with", "about", "against", 
        "between", "into", "through", "during", "before", "after", "above", "below", "up", "down", 
        "on", "off", "again", "further", "once", "here", "there", "all", "both", "each", "few", 
        "most", "other", "some", "such", "nor", "not", "only", "own", "same", "so", "too", "very", 
        "can", "will", "just", "should", "now", "sewak", "singh"
    }

    import re
    # Split text into tokens (words and non-words) to preserve layout exactly
    tokens = re.split(r'(\s+|[^\w\s\-\/\.]+)', text)
    
    result = []
    for token in tokens:
        if not token or token.isspace():
            result.append(token)
            continue
            
        # If token is already Gurmukhi Unicode, keep it
        if any('\u0A00' <= c <= '\u0A7F' for c in token):
            result.append(token)
            continue
            
        # Clean token for lookup
        clean_token = token.strip('.,()-"\'/:').lower()
        
        # If it's a number, keep it
        if clean_token.isdigit() or re.match(r'^\d+[\.\d]*$', clean_token):
            result.append(token)
            continue
            
        # If it is a known English word, keep it
        if clean_token in english_words:
            result.append(token)
        else:
            # Convert legacy to Unicode
            converted_token = conv.convert_font(token, from_font, "Unicode")
            result.append(converted_token)
            
    return "".join(result)

@app.route('/convert', methods=['POST'])
def convert_text():
    data = request.get_json()
    text = data.get('text', '')
    from_font = data.get('from_font', 'auto')
    to_font = data.get('to_font', 'Unicode')

    if not text:
        return jsonify({'converted': ''})

    # Use the same logic as desktop: if from_font is "auto", detect
    if from_font == "auto":
        if is_gurmukhi_unicode(text):
            from_font = CONV_SOURCE_UNICODE  # "AnmolUni"
        else:
            from_font = detect_legacy_font(text) or "AnmolLipi"

    try:
        import html
        converted_html = ""
        
        if from_font == CONV_SOURCE_UNICODE and to_font == CONV_SOURCE_UNICODE:
            converted = text
            converted_html = html.escape(text)
        else:
            # Case A: Unicode to Legacy
            if from_font in ("Unicode", "AnmolUni") and to_font != "Unicode":
                segments = split_gurmukhi_segments(text)
                converted_parts = []
                html_parts = []
                for is_gur, seg_text in segments:
                    if is_gur:
                        conv_text = conv.convert_font(seg_text, from_font, to_font)
                        converted_parts.append(conv_text)
                        html_parts.append(f"<span class=\"font-legacy\">{html.escape(conv_text)}</span>")
                    else:
                        converted_parts.append(seg_text)
                        html_parts.append(f"<span class=\"font-english\">{html.escape(seg_text)}</span>")
                converted = "".join(converted_parts)
                converted_html = "".join(html_parts)
            # Case B: Legacy to Unicode (Using our robust mixed word converter)
            elif to_font == "Unicode":
                converted = convert_mixed_legacy_to_unicode(text, from_font, conv)
                converted_html = html.escape(converted)
            # Case C: Legacy to Legacy
            else:
                converted = conv.convert_font(text, from_font, to_font)
                converted_html = f"<span class=\"font-legacy\">{html.escape(converted)}</span>"
                
        return jsonify({
            'converted': sanitize_unicode(converted),
            'converted_html': converted_html
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

def convert_docx_run(run, paragraph, from_font, to_font, conv):
    """
    Convert a single docx run, splitting it into Gurmukhi and non-Gurmukhi segments
    if converting to legacy, applying appropriate fonts and preserving original formatting.
    """
    ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    text = run.text
    if not text or not text.strip():
        return
        
    # Get run's font name from run properties XML or python-docx object
    font_name = run.font.name
    if not font_name:
        rPr = run._r.find(f'{ns}rPr')
        if rPr is not None:
            rFonts = rPr.find(f'{ns}rFonts')
            if rFonts is not None:
                font_name = rFonts.get(f'{ns}ascii') or rFonts.get(f'{ns}hAnsi')
                
    # Fallback to paragraph style font name if inherited
    if not font_name and paragraph.style and paragraph.style.font:
        font_name = paragraph.style.font.name

    # --- Case 1: Unicode to Legacy ---
    if to_font != "Unicode":
        # Only convert if text has Gurmukhi characters
        if is_gurmukhi_unicode(text):
            # Split run into Gurmukhi and English segments
            segments = split_gurmukhi_segments(text)
            if not segments:
                return
                
            # If all segments are non-Gurmukhi, leave as is
            if all(not is_gur for is_gur, _ in segments):
                return
                
            r_element = run._r
            parent = r_element.getparent()
            previous_element = r_element
            
            for is_gur, seg_text in segments:
                if not seg_text:
                    continue
                new_r = copy.deepcopy(r_element)
                t_el = new_r.find(f'{ns}t')
                if t_el is not None:
                    if is_gur:
                        # Convert Gurmukhi segment
                        t_el.text = conv.convert_font(seg_text, CONV_SOURCE_UNICODE, to_font)
                        
                        # Apply legacy font family
                        rPr = new_r.find(f'{ns}rPr')
                        if rPr is None:
                            rPr = OxmlElement('w:rPr')
                            new_r.insert(0, rPr)
                        rFonts = rPr.find(f'{ns}rFonts')
                        if rFonts is None:
                            rFonts = OxmlElement('w:rFonts')
                            rPr.append(rFonts)
                        
                        legacy_font = WORD_FONT_MAP.get(to_font, "Asees")
                        rFonts.set(f'{ns}ascii', legacy_font)
                        rFonts.set(f'{ns}hAnsi', legacy_font)
                    else:
                        t_el.text = seg_text
                        # For English segment, preserve formatting and original font
                
                parent.insert(parent.index(previous_element) + 1, new_r)
                previous_element = new_r
            parent.remove(r_element)

    # --- Case 2: Legacy to Unicode ---
    else:
        # Detect legacy source font
        src = from_font
        if src == "auto":
            if font_name in LEGACY_FONTS:
                src = font_name
            else:
                # If font name is not explicit, use fallback detection
                src = detect_legacy_font(text)
                
        # Only convert if the run is explicitly styled with legacy font or detected as legacy
        is_legacy = (font_name in LEGACY_FONTS) or (not font_name and detect_legacy_font(text))
        
        if is_legacy and src and src != "Unicode":
            # Convert text
            converted_text = conv.convert_font(text, src, CONV_SOURCE_UNICODE)
            run.text = converted_text
            
            # Clear legacy font name or change to Unicode font (Raavi)
            rPr = run._r.find(f'{ns}rPr')
            if rPr is not None:
                rFonts = rPr.find(f'{ns}rFonts')
                if rFonts is not None:
                    rFonts.set(f'{ns}ascii', 'Raavi')
                    rFonts.set(f'{ns}hAnsi', 'Raavi')

@app.route('/convert_docx', methods=['POST'])
def convert_docx():
    data = request.get_json()
    file_b64 = data.get('file_b64')
    from_font = data.get('from_font', 'auto')
    to_font = data.get('to_font', 'Unicode')
    if not file_b64:
        return jsonify({'error': 'No file data'}), 400

    try:
        docx_bytes = base64.b64decode(file_b64)
        doc = Document(io.BytesIO(docx_bytes))
        
        # Process paragraphs
        for paragraph in doc.paragraphs:
            runs = list(paragraph.runs)
            for run in runs:
                convert_docx_run(run, paragraph, from_font, to_font, conv)
                
        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        runs = list(paragraph.runs)
                        for run in runs:
                            convert_docx_run(run, paragraph, from_font, to_font, conv)
                            
        out = io.BytesIO()
        doc.save(out)
        out.seek(0)
        out_b64 = base64.b64encode(out.read()).decode('utf-8')
        return jsonify({'file_b64': out_b64})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

def translate_via_gemini(text: str, src_lang: str, dest_lang: str, api_key: str) -> str:
    url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash:generateContent?key={api_key}"
    
    prompt = f"""
    You are an expert legal translator specializing in Indian court, FIR, and police documents.
    Translate the following text from {src_lang} to {dest_lang}.
    
    CRITICAL TRANSLATION GUIDELINES:
    1. Use formal, professional, and precise Indian legal, court, and police terminology.
       - Translate 'ਸੇਵਾ ਵਿਖੇ' / 'seva vikhe' as 'To,'
       - Translate 'ਸੀਨੀਅਰ ਕਪਤਾਨ ਪੁਲਿਸ' as 'Senior Superintendent of Police (SSP)'
       - Translate 'ਸ੍ਰੀਮਾਨ ਜੀ' / 'ਸ਼੍ਰੀਮਾਨ ਜੀ' / 'ਸ੍ਰੀਮਾਨ ਜੀ,' as 'Respected Sir,'
       - Translate 'ਪੁਲਿਸ ਲਾਇਨ' / 'ਪੁਲਿਸ ਲਾਇਨਜ਼' as 'Police Lines'
       - Translate 'ਅਗਲੇਰੀ ਯੋਗ ਕਾਰਵਾਈ ਲਈ ਪੇਸ਼ ਹੈ ਜੀ' as 'Submitted for further necessary action, please.'
       - Translate 'ਇੰਸਪੈਕਟਰ ਪੁਲਿਸ' as 'Inspector of Police'
       - Translate 'ਇੰਚਾਰਜ' as 'Incharge'
       - Translate 'ਟੈਕਨੀਕਲ ਸਰਵਿਸਿਜ਼' as 'Technical Services'
       - Translate 'Qwxw' / 'Qwxy' / 'ਥਾਣਾ' as 'Police Station (P.S.)'
       - Translate 'Sikwieqkrqw' / 'sUcnwkrqw' / 'ਸ਼ਿਕਾਇਤਕਰਤਾ' / 'ਸੂਚਨਾਕਰਤਾ' as 'Complainant/Informant'
       - Translate 'mulzm' / 'doSI' / 'ਮੁਲਜ਼ਮ' / 'ਦੋਸ਼ੀ' as 'Accused'
       - Translate 'Dwrw' / 'ਧਾਰਾ' as 'Section'
       - Translate 'AprwD' / 'ਅਪਰਾਧ' as 'Offence'
       - Translate 'zilHw' / 'ਜ਼ਿਲ੍ਹਾ' as 'District'
       - Translate 'rojnwmcw' / 'ਰੋਜ਼ਨਾਮਚਾ' as 'General Diary (GD)'
       - Translate 'lVI' / 'lVI nM.' / 'ਲੜੀ ਨੰਬਰ' as 'S.No.' / 'Serial Number'
       - Translate 'ਪੁਲਿਸ ਪਾਰਟੀ' as 'Police Party'
       - Translate 'ਗਸ਼ਤ' as 'Patrolling'
       - Translate 'ਨਾਕਾਬੰਦੀ' as 'Naka Bandi (Checking)'
       - Translate 'ਤਫਤੀਸ਼ੀ ਅਫ਼ਸਰ' / 'investigating officer' as 'Investigating Officer (I.O.)'
       - Translate 'ਰਸਤਾ' / 'ਸੜਕ' as 'Route / Road'
       - Translate 'ਮੌਕਾ' as 'Place of Occurrence / Spot'
       - Translate 'ਬਰਾਮਦਗੀ' as 'Recovery'
       - Translate 'ਰੂਪੋਸ਼' as 'Absconding'
       - Translate 'ਹਿਰਾਸਤ' as 'Custody'
       - Translate 'ਜਾਮਾ ਤਲਾਸ਼ੀ' as 'Personal Search'
       - Translate 'ਸਰਕਾਰੀ ਗੱਡੀ' as 'Government Vehicle'
       - Translate BNS / BNSS / NDPS terms formally and accurately.
    2. Maintain the layout, line breaks, structures, numbers, dates, times, and punctuation exactly as in the original text.
    3. Do not omit any details, names, dates, or numbers.
    4. Provide only the translated text. Do not include any notes, explanations, or conversational preamble.
    
    Original Text:
    {text}
    """
    
    payload = {
        "contents": [
            {
                "parts": [
                    {"text": prompt}
                ]
            }
        ]
    }
    
    req_data = json.dumps(payload).encode('utf-8')
    req = urllib.request.Request(
        url,
        data=req_data,
        headers={'Content-Type': 'application/json'}
    )
    
    try:
        with urllib.request.urlopen(req, timeout=30) as response:
            res_body = response.read().decode('utf-8')
            res_json = json.loads(res_body)
            text = res_json['candidates'][0]['content']['parts'][0]['text']
            return text
    except urllib.error.HTTPError as e:
        error_msg = e.read().decode('utf-8')
        try:
            err_json = json.loads(error_msg)
            message = err_json.get('error', {}).get('message', str(e))
        except:
            message = error_msg or str(e)
        raise Exception(f"Gemini Translation Error: {message}")
    except Exception as e:
        raise Exception(f"Failed to connect to Gemini: {str(e)}")

def translate_via_google_free(text: str, src: str, dest: str) -> str:
    if not text.strip():
        return text
        
    # Split into chunks of max 4000 characters, keeping paragraph structure intact
    paragraphs = text.split('\n')
    chunks = []
    current_chunk = []
    current_len = 0
    
    for p in paragraphs:
        if current_len + len(p) + 1 > 4000:
            chunks.append("\n".join(current_chunk))
            current_chunk = [p]
            current_len = len(p)
        else:
            current_chunk.append(p)
            current_len += len(p) + 1
            
    if current_chunk:
        chunks.append("\n".join(current_chunk))
        
    translated_chunks = []
    for chunk in chunks:
        if not chunk.strip():
            translated_chunks.append(chunk)
            continue
            
        # URL encode the query text
        query = urllib.parse.quote(chunk)
        url = f"https://translate.googleapis.com/translate_a/single?client=gtx&sl={src}&tl={dest}&dt=t&q={query}"
        
        req = urllib.request.Request(
            url,
            headers={'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64)'}
        )
        
        try:
            with urllib.request.urlopen(req, timeout=15) as response:
                res_body = response.read().decode('utf-8')
                res_json = json.loads(res_body)
                sentences = res_json[0]
                translated_text = "".join(sentence[0] for sentence in sentences if sentence[0])
                translated_chunks.append(translated_text)
        except Exception as e:
            print(f"Google free translation failed for chunk: {e}")
            try:
                translated_text = GoogleTranslator(source=src, target=dest).translate(chunk)
                translated_chunks.append(translated_text)
            except:
                translated_chunks.append(chunk)
                
    return "\n".join(translated_chunks)

def post_process_legal_translation(text: str) -> str:
    # A dictionary of case-insensitive replacements to formalize Indian police/court terminology
    replacements = {
        r'\bat the service\b': 'To,',
        r'\bin service\b': 'To,',
        r'\bsenior captain police\b': 'Senior Superintendent of Police (SSP)',
        r'\bsenior captain of police\b': 'Senior Superintendent of Police (SSP)',
        r'\bahs g\b': 'Respected Sir,',
        r'\bshriman ji\b': 'Respected Sir,',
        r'\bsriman ji\b': 'Respected Sir,',
        r'\bpolice line\b': 'Police Lines',
        r'\bpolice lines\b': 'Police Lines',
        r'\bdistrict training school\b': 'District Training School (D.T.S.)',
        r'\bsubmitted for next eligible action\b': 'Submitted for further necessary action, please.',
        r'\bpresented for further action\b': 'Submitted for further necessary action, please.',
        r'\bpresented for further necessary action\b': 'Submitted for further necessary action, please.',
        r'\binspector of police\b': 'Inspector of Police',
        r'\btechnical services\b': 'Technical Services',
        r'\bin charge\b': 'Incharge',
        r'\bin-charge\b': 'Incharge',
        r'\bwith thanks\b': 'With thanks',
        r'\bpolice station\b': 'Police Station (P.S.)',
        r'\bpolice stations\b': 'Police Stations (P.S.)',
        r'\bcomplainant\b': 'Complainant/Informant',
        r'\binformant\b': 'Complainant/Informant',
        r'\baccused\b': 'Accused',
        r'\bsection\b': 'Section',
        r'\bsections\b': 'Sections',
        r'\boffence\b': 'Offence',
        r'\boffences\b': 'Offences',
        r'\bdistrict\b': 'District',
        r'\bgeneral diary\b': 'General Diary (GD)',
        r'\bserial number\b': 'S.No. / Serial Number',
        r'\bs\.no\b': 'S.No.',
        r'\bpolice party\b': 'Police Party',
        r'\bpatrolling\b': 'Patrolling',
        r'\bnaka bandi\b': 'Naka Bandi (Checking)',
        r'\binvestigating officer\b': 'Investigating Officer (I.O.)',
        r'\bplace of occurrence\b': 'Place of Occurrence / Spot',
        r'\brecovery\b': 'Recovery',
        r'\babsconding\b': 'Absconding',
        r'\bcustody\b': 'Custody',
        r'\bpersonal search\b': 'Personal Search',
        r'\bgovernment vehicle\b': 'Government Vehicle'
    }
    
    import re
    processed = text
    for pattern, replacement in replacements.items():
        processed = re.sub(pattern, replacement, processed, flags=re.IGNORECASE)
    return processed

@app.route('/translate', methods=['POST'])
def translate():
    data = request.get_json()
    text = data.get('text', '')
    src = data.get('src', 'auto')
    dest = data.get('dest', 'en')
    
    # Check if header contains Gemini Key
    gemini_key = request.headers.get('X-Gemini-API-Key')
    
    if not text:
        return jsonify({'translation': ''})
        
    try:
        # Step 1: Detect and handle legacy text input
        # If the input text is legacy ASCII (not Unicode), we must convert it to Unicode Punjabi first
        # so that translation models can parse it.
        if not is_gurmukhi_unicode(text):
            detected = detect_legacy_font(text)
            if detected and detected != "Unicode":
                text = convert_mixed_legacy_to_unicode(text, detected, conv)
        
        # Step 2: Perform translation
        if gemini_key:
            # Map abbreviations to full names for Gemini prompt context
            lang_names = {
                'pa': 'Punjabi',
                'hi': 'Hindi',
                'en': 'English',
                'ur': 'Urdu',
                'auto': 'Auto Detect'
            }
            src_name = lang_names.get(src, 'Auto-Detected Language')
            dest_name = lang_names.get(dest, 'English')
            
            translated = translate_via_gemini(text, src_name, dest_name, gemini_key)
        else:
            raw_translation = translate_via_google_free(text, src, dest)
            translated = post_process_legal_translation(raw_translation)
            
        return jsonify({'translation': translated})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True)